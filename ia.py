from tkinter import *
from tkinter import messagebox
from tkinter import filedialog
import sqlite3
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configurações de cores
co0 = "#2f2f2f"  # Dark grey
co1 = "#3b3f4e"  # Dark blue-grey
co2 = "#33cc33"  # Bright green (for highlighting words)
co3 = "#4b5154"  # Dark grey-brown
co4 = "#ffffff"  # White

# Conectar ao banco de dados (ou criar se não existir)
conexao = sqlite3.connect('frases.db')
cursor = conexao.cursor()

# Criar a tabela de frases (adicionar a coluna 'cluster' se não existir)
cursor.execute('''
CREATE TABLE IF NOT EXISTS frases (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    cluster TEXT,
    frase TEXT NOT NULL
)
''')

# Adicionar coluna 'cluster' se não existir
try:
    cursor.execute('ALTER TABLE frases ADD COLUMN cluster TEXT')
except sqlite3.OperationalError:
    # Se a coluna já existir, ignorar o erro
    pass

# Variável global para a lista de frases
lista_frases = None

# Função para mostrar a janela de gerenciamento de frases
def mostrar_janela_gerenciamento():
    global lista_frases

    janela_login.destroy()  # Fechar a janela de login

    root = Tk()
    root.title("Gerenciador de Frases")
    root.configure(background=co1)

    # Frame para inserir e visualizar frases
    frame_inserir_visualizar = Frame(root, bg=co1)
    frame_inserir_visualizar.pack(padx=10, pady=10)

    Label(frame_inserir_visualizar, text="Selecione o cluster:", bg=co1, fg=co4).grid(row=0, column=0, padx=5, pady=5)
    variavel_cluster = StringVar(root)
    variavel_cluster.set("Cluster 1")  # Valor padrão
    clusters = ["Cluster 1", "Cluster 2", "Cluster 3", "Cluster 4", "Cluster 5"]
    menu_cluster = OptionMenu(frame_inserir_visualizar, variavel_cluster, *clusters)
    menu_cluster.config(bg=co1, fg=co4)
    menu_cluster.grid(row=0, column=1, padx=5, pady=5)

    Label(frame_inserir_visualizar, text="Digite uma nova frase:", bg=co1, fg=co4).grid(row=1, column=0, padx=5, pady=5)
    entrada_frase = Entry(frame_inserir_visualizar, width=50, highlightthickness=1, relief="solid")
    entrada_frase.grid(row=1, column=1, padx=5, pady=5)

    botao_inserir = Button(frame_inserir_visualizar, text="Inserir Frase", command=lambda: inserir_frase(variavel_cluster.get(), entrada_frase), bg=co2, fg=co1)
    botao_inserir.grid(row=1, column=2, padx=5, pady=5)

    Label(frame_inserir_visualizar, text="ID da frase a ser deletada:", bg=co1, fg=co4).grid(row=2, column=0, padx=5, pady=5)
    entrada_id = Entry(frame_inserir_visualizar, width=10, highlightthickness=1, relief="solid")
    entrada_id.grid(row=2, column=1, padx=5, pady=5)

    botao_deletar = Button(frame_inserir_visualizar, text="Deletar Frase", command=lambda: deletar_frase(entrada_id), bg=co2, fg=co1)
    botao_deletar.grid(row=2, column=2, padx=5, pady=5)

    botao_visualizar = Button(frame_inserir_visualizar, text="Visualizar Frases", command=visualizar_frases, bg=co2, fg=co1)
    botao_visualizar.grid(row=3, column=1, padx=5, pady=5)

    # Botão para exportar frases para o Word
    botao_exportar_word = Button(frame_inserir_visualizar, text="Exportar Frases para Word", command=exportar_para_word, bg=co2, fg=co1)
    botao_exportar_word.grid(row=4, column=1, padx=5, pady=5)

    # Lista de frases
    frame_lista = Frame(root, bg=co1)
    frame_lista.pack(padx=10, pady=10)

    Label(frame_lista, text="Frases Armazenadas:", bg=co1, fg=co4).pack()

    scrollbar = Scrollbar(frame_lista)
    scrollbar.pack(side=RIGHT, fill=Y)

    lista_frases = Listbox(frame_lista, width=70, yscrollcommand=scrollbar.set)
    lista_frases.pack()

    scrollbar.config(command=lista_frases.yview)

    # Executar a interface
    root.mainloop()

def inserir_frase(cluster, entrada_frase):
    frase = entrada_frase.get()
    if cluster.strip() == "" or frase.strip() == "":
        messagebox.showerror("Erro", "Selecione um cluster e digite uma frase válida!")
        return

    cursor.execute('INSERT INTO frases (cluster, frase) VALUES (?, ?)', (cluster, frase))
    conexao.commit()
    entrada_frase.delete(0, END)  # Limpar os campos após a inserção

def visualizar_frases():
    global lista_frases
    cursor.execute('SELECT * FROM frases')
    frases = cursor.fetchall()
    if not frases:
        messagebox.showinfo("Aviso", "Não há frases armazenadas!")
    else:
        lista_frases.delete(0, END)
        for frase in frases:
            lista_frases.insert(END, f"{frase[0]}: {frase[1]} - {frase[2]}")

def deletar_frase(entrada_id):
    id_frase = entrada_id.get()
    if id_frase.strip() == "":
        messagebox.showerror("Erro", "Digite o ID da frase a ser deletada!")
        return

    try:
        cursor.execute('DELETE FROM frases WHERE id = ?', (id_frase,))
        conexao.commit()
        entrada_id.delete(0, END)  # Limpar o campo após a deleção
    except sqlite3.Error as e:
        messagebox.showerror("Erro", f"Erro ao deletar a frase: {e}")

def exportar_para_word():
    nome_arquivo = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Arquivo do Word", "*.docx")])
    if not nome_arquivo:
        return

    cursor.execute('SELECT * FROM frases')
    frases = cursor.fetchall()

    documento = Document()
    documento.add_heading('Segmentos de Textos', 0)

    # Criar tabela com 2 colunas
    tabela = documento.add_table(rows=1, cols=2)
    tabela.style = 'Table Grid'  # Estilo da tabela com bordas

    # Configurar largura das colunas
    largura_coluna_cluster = Inches(1.0)
    largura_coluna_frase = Inches(4.0)
    tabela.columns[0].width = largura_coluna_cluster
    tabela.columns[1].width = largura_coluna_frase

    # Adicionar cabeçalho à tabela
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Cluster'
    hdr_cells[1].text = 'Segmento de Texto'
    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionar frases à tabela
    for frase in frases:
        row_cells = tabela.add_row().cells
        row_cells[0].text = frase[1]  # Cluster
        row_cells[1].text = frase[2]  # Frase
        row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Salvar documento
    try:
        documento.save(nome_arquivo)
        messagebox.showinfo("Sucesso", f"Tabela exportada para {nome_arquivo}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao exportar a tabela: {e}")

# Função para validar o login
def validar_login():
    usuario = entrada_usuario.get()
    senha = entrada_senha.get()
    if usuario == "jppacheco" and senha == "1512":
        mostrar_janela_gerenciamento()
    else:
        messagebox.showerror("Erro", "Usuário ou senha incorretos")

# Interface de login
janela_login = Tk()
janela_login.title("")
janela_login.geometry('310x300')
janela_login.configure(background=co1)
janela_login.resizable(width=False, height=False)

################# Frames ####################

frame_cima = Frame(janela_login, width=310, height=50, bg=co1, relief="flat")
frame_cima.grid(row=0, column=0, pady=1, padx=0, sticky=NSEW)

frame_baixo = Frame(janela_login, width=310, height=400, bg=co1, relief="flat")
frame_baixo.grid(row=1, column=0, pady=1, padx=0, sticky=NSEW)

# Configurando frame_cima

l_nome = Label(frame_cima, text="LOGIN", height=6, anchor=NE, font=('Ivy 25 '), bg=co1, fg=co4)
l_nome.place(x=5, y=5)

l_linha = Label(frame_cima, width=275, text="", height=1, anchor=NW, font=('Ivy 1 '), bg=co2)
l_linha.place(x=10, y=45)

# Configurando frame_baixo ---------------------------

l_nome = Label(frame_baixo, text="Nome *", height=1, anchor=NW, font=('Ivy 10 bold'), bg=co1, fg=co4)
l_nome.place(x=10, y=20)
entrada_usuario = Entry(frame_baixo, width=25, justify='left', font=("", 15), highlightthickness=1, relief="solid")
entrada_usuario.place(x=14, y=50)

l_pass = Label(frame_baixo, text="Senha *", height=1, anchor=NW, font=('Ivy 10 bold'), bg=co1, fg=co4)
l_pass.place(x=10, y=95)
entrada_senha = Entry(frame_baixo, show='*', width=25, justify='left', font=("", 15), highlightthickness=1, relief="solid")
entrada_senha.place(x=15, y=130)

botao_confirmar = Button(frame_baixo, text="Entrar", width=39, height=2, bg=co2, fg=co1, font=('Ivy 8 bold'), relief=RAISED, overrelief=RIDGE, command=validar_login)
botao_confirmar.place(x=15, y=180)

janela_login.mainloop()

# Fechar a conexão com o banco de dados ao finalizar o programa
conexao.close()
